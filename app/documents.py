"""Генерация документов"""

import io
from datetime import datetime

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side


def generate_act_docx(ticket):
    
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("АКТ ВЫПОЛНЕННЫХ РАБОТ")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph(f"по заявке № {ticket.number}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        f"от {datetime.utcnow().strftime('%d.%m.%Y')}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    doc.add_paragraph(
        "Служба технической поддержки университета выполнила следующие работы:"
    )

    rows = [
        ("Номер заявки", ticket.number),
        ("Тема", ticket.title),
        ("Категория", ticket.category.name if ticket.category else "—"),
        ("Приоритет", ticket.priority.name if ticket.priority else "—"),
        ("Заявитель", ticket.author.full_name if ticket.author else "—"),
        ("Исполнитель", ticket.assignee.full_name if ticket.assignee else "—"),
        ("Статус", ticket.status.name if ticket.status else "—"),
        ("Дата создания", ticket.created_at.strftime("%d.%m.%Y %H:%M") if ticket.created_at else "—"),
        ("Дата закрытия", ticket.closed_at.strftime("%d.%m.%Y %H:%M") if ticket.closed_at else "—"),
    ]

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = str(value)
        cells[0].paragraphs[0].runs[0].bold = True if cells[0].paragraphs[0].runs else None

    doc.add_paragraph("")
    doc.add_paragraph("Описание проблемы:").runs[0].bold = True
    doc.add_paragraph(ticket.description)

    doc.add_paragraph("")
    doc.add_paragraph("Исполнитель: _______________ / ______________________")
    doc.add_paragraph("Заявитель:   _______________ / ______________________")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_report_xlsx(tickets):
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Отчет по заявкам"

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="2F5496")
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin = Side(style="thin", color="999999")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    title = ws.cell(row=1, column=1, value="Отчет по заявкам службы технической поддержки")
    title.font = Font(bold=True, size=14)
    ws.merge_cells("A1:H1")
    ws.cell(row=2, column=1,
            value=f"Сформирован: {datetime.utcnow().strftime('%d.%m.%Y %H:%M')}")
    ws.merge_cells("A2:H2")

    headers = ["№", "Номер", "Тема", "Категория", "Приоритет",
               "Статус", "Заявитель", "Дата создания"]
    for col, h in enumerate(headers, start=1):
        c = ws.cell(row=4, column=col, value=h)
        c.font = header_font
        c.fill = header_fill
        c.alignment = center
        c.border = border

    for i, t in enumerate(tickets, start=1):
        row = 4 + i
        values = [
            i, t.number, t.title,
            t.category.name if t.category else "—",
            t.priority.name if t.priority else "—",
            t.status.name if t.status else "—",
            t.author.full_name if t.author else "—",
            t.created_at.strftime("%d.%m.%Y") if t.created_at else "—",
        ]
        for col, v in enumerate(values, start=1):
            c = ws.cell(row=row, column=col, value=v)
            c.border = border
            c.alignment = Alignment(vertical="center", wrap_text=True)

    widths = [5, 12, 35, 20, 14, 14, 25, 14]
    for col, w in enumerate(widths, start=1):
        ws.column_dimensions[ws.cell(row=4, column=col).column_letter].width = w

    total_row = 4 + len(tickets) + 1
    tc = ws.cell(row=total_row, column=1, value=f"Всего заявок: {len(tickets)}")
    tc.font = Font(bold=True)
    ws.merge_cells(start_row=total_row, start_column=1, end_row=total_row, end_column=8)

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer
